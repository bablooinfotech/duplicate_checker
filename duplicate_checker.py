import pandas as pd
from pathlib import Path
from docx import Document
import os
from dotenv import load_dotenv

load_dotenv()
# ==========================================================
# READ EXCEL FILE
# ==========================================================

# def read_excel_file(file_path):
#     try:
#         print(f"\nReading Excel: {file_path.name}")

#         excel = pd.read_excel(
#             file_path,
#             header=None,
#             dtype=str
#         )

#         header_row = None

#         for idx, row in excel.iterrows():

#             values = row.astype(str).tolist()

#             non_empty_count = sum(
#                 1
#                 for val in values
#                 if str(val).strip() not in ["", "nan"]
#             )

#             if non_empty_count >= 5:
#                 header_row = idx
#                 break

#         if header_row is None:
#             print(
#                 f"Header not found in {file_path.name}"
#             )
#             return pd.DataFrame()

#         df = pd.read_excel(
#             file_path,
#             header=header_row,
#             dtype=str
#         )

#         df.columns = [
#             str(col).strip()
#             for col in df.columns
#         ]

#         df = df.dropna(how="all")

#         print(
#             f"Loaded {len(df)} records from "
#             f"{file_path.name}"
#         )

#         return df

#     except Exception as e:
#         print(
#             f"Error reading {file_path.name}: {e}"
#         )
#         return pd.DataFrame()

def read_excel_file(file_path):
    try:
        print(f"\nReading Excel: {file_path.name}")

        if file_path.suffix.lower() == ".xls":
            engine = "xlrd"
        else:
            engine = "openpyxl"

        excel = pd.read_excel(
            file_path,
            header=None,
            dtype=str,
            engine=engine
        )

        header_row = None

        for idx, row in excel.iterrows():

            values = row.astype(str).tolist()

            non_empty_count = sum(
                1
                for val in values
                if str(val).strip() not in ["", "nan"]
            )

            if non_empty_count >= 5:
                header_row = idx
                break

        if header_row is None:
            print(
                f"Header not found in {file_path.name}"
            )
            return pd.DataFrame()

        df = pd.read_excel(
            file_path,
            header=header_row,
            dtype=str,
            engine=engine
        )

        df.columns = [
            str(col).strip()
            for col in df.columns
        ]

        df = df.dropna(how="all")

        print(
            f"Loaded {len(df)} records from "
            f"{file_path.name}"
        )

        return df

    except Exception as e:
        print(
            f"Error reading {file_path.name}: {e}"
        )
        return pd.DataFrame()

# ==========================================================
# READ DOCX FILE
# ==========================================================

def read_docx_file(file_path):
    try:
        print(f"\nReading DOCX: {file_path.name}")

        document = Document(file_path)

        data = []

        EXPECTED_HEADERS = {
            "company",
            "vacancy",
            "post",
            "qualification",
            "salary",
            "job location",
            "employer name",
            "mobile"
        }

        for table in document.tables:

            if len(table.rows) < 2:
                continue

            header_row_index = None
            headers = None

            # Find actual header row
            for idx, row in enumerate(table.rows):

                row_values = [
                    str(cell.text).strip()
                    for cell in row.cells
                ]

                row_text = " ".join(
                    row_values
                ).lower()

                matched_headers = sum(
                    1
                    for header in EXPECTED_HEADERS
                    if header in row_text
                )

                if matched_headers >= 3:
                    header_row_index = idx
                    headers = row_values
                    break

            if header_row_index is None:
                continue

            headers = [
                str(col).strip()
                for col in headers
            ]

            # Remove empty headers
            headers = [
                h if h else f"column_{i}"
                for i, h in enumerate(headers)
            ]

            # Read data rows
            for row in table.rows[header_row_index + 1:]:

                values = [
                    str(cell.text).strip()
                    for cell in row.cells
                ]

                # Pad missing values
                if len(values) < len(headers):
                    values.extend(
                        [""] * (
                            len(headers) - len(values)
                        )
                    )

                # Trim extra values
                values = values[:len(headers)]

                # Skip completely blank rows
                if not any(values):
                    continue

                data.append(
                    dict(zip(headers, values))
                )

        df = pd.DataFrame(data)

        # Remove duplicate columns if any
        df = df.loc[:, ~df.columns.duplicated()]

        print(
            f"Loaded {len(df)} records from "
            f"{file_path.name}"
        )

        return df

    except Exception as e:
        print(
            f"Error reading {file_path.name}: {e}"
        )
        return pd.DataFrame()
    
# ==========================================================
# READ CSV FILE
# ==========================================================

def read_docx_file(file_path):
    try:
        print(f"\nReading DOCX: {file_path.name}")

        document = Document(file_path)

        all_data = []

        EXPECTED_HEADERS = {
            "sl no",
            "company",
            "vacancy",
            "post",
            "qualification",
            "salary",
            "job location",
            "employer name",
            "mobile"
        }

        for table_index, table in enumerate(document.tables):

            if len(table.rows) < 2:
                continue

            headers = None

            # Find header row
            for row in table.rows:

                row_values = [
                    str(cell.text).strip()
                    for cell in row.cells
                ]

                row_text = " ".join(
                    row_values
                ).lower()

                match_count = sum(
                    1
                    for header in EXPECTED_HEADERS
                    if header in row_text
                )

                if match_count >= 3:
                    headers = row_values
                    break

            if not headers:
                continue

            headers = [
                str(h).strip()
                for h in headers
            ]

            print(
                f"Table {table_index + 1}: "
                f"Detected {len(headers)} columns"
            )

            for row in table.rows:

                values = [
                    str(cell.text).strip()
                    for cell in row.cells
                ]

                if len(values) != len(headers):
                    continue

                # Skip header repetitions
                row_text = " ".join(
                    values
                ).lower()

                header_match = sum(
                    1
                    for header in EXPECTED_HEADERS
                    if header in row_text
                )

                if header_match >= 3:
                    continue

                # Skip empty rows
                if not any(
                    value.strip()
                    for value in values
                ):
                    continue

                record = dict(
                    zip(headers, values)
                )

                all_data.append(record)

        df = pd.DataFrame(all_data)

        if not df.empty:

            df.columns = [
                str(col).strip()
                for col in df.columns
            ]

            df = df.loc[
                :,
                ~df.columns.duplicated()
            ]

            df = df.dropna(
                how="all"
            )

            df = df.reset_index(
                drop=True
            )

        print(
            f"Loaded {len(df)} records from "
            f"{file_path.name}"
        )

        return df

    except Exception as e:

        print(
            f"Error reading {file_path.name}: {e}"
        )

        return pd.DataFrame()
# ==========================================================
# LOAD FILES
# ==========================================================

def load_all_files(folder_path):

    file_data = {}

    files = list(Path(folder_path).glob("*"))

    print(f"\nFound {len(files)} files")

    for file in files:

        df = pd.DataFrame()

        if file.suffix.lower() in [".xlsx", ".xls"]:
            df = read_excel_file(file)

        elif file.suffix.lower() == ".docx":
            df = read_docx_file(file)

        elif file.suffix.lower() == ".csv":
            df = read_csv_file(file)

        if not df.empty:
            file_data[file.stem] = df

    return file_data


# ==========================================================
# VALIDATE COLUMNS
# ==========================================================

def validate_columns(df, columns):

    missing = [
        col
        for col in columns
        if col not in df.columns
    ]

    if missing:
        raise ValueError(
            f"Columns not found: {missing}"
        )


# ==========================================================
# DUPLICATE DETECTION
# ==========================================================

def detect_duplicates(df, columns):

    duplicates_df = df[
        df.duplicated(
            subset=columns,
            keep=False
        )
    ].copy()

    unique_df = df.drop_duplicates(
        subset=columns,
        keep="first"
    ).copy()

    return duplicates_df, unique_df


# ==========================================================
# EXPORT FILES
# ==========================================================

def export_results(
    output_folder,
    file_name,
    duplicates_df,
    unique_df
):

    duplicate_file = (
        Path(output_folder)
        / f"{file_name}_duplicates.xlsx"
    )

    unique_file = (
        Path(output_folder)
        / f"{file_name}_unique.xlsx"
    )

    duplicates_df.to_excel(
        duplicate_file,
        index=False
    )

    unique_df.to_excel(
        unique_file,
        index=False
    )

    print(
        f"Created: {duplicate_file.name}"
    )

    print(
        f"Created: {unique_file.name}"
    )


# ==========================================================
# MAIN
# ==========================================================

def main():

    print("\n================================")
    print(" DUPLICATE DATA CHECKER ")
    print("================================")

    # input_folder = (
    #     r"C:\Users\REAL ENTERPRISE\OneDrive\Desktop"
    #     r"\learning\duplicate_checker\input_files"
    # )

    # output_folder = (
    #     r"C:\Users\REAL ENTERPRISE\OneDrive\Desktop"
    #     r"\learning\duplicate_checker\output"
    # )
    
    input_folder = os.getenv("INPUT_FOLDER")

    output_folder = os.getenv("OUTPUT_FOLDER")

    os.makedirs(
        output_folder,
        exist_ok=True
    )

    file_data = load_all_files(
        input_folder
    )

    if not file_data:
        print("\nNo valid files found.")
        return

    first_df = next(
        iter(file_data.values())
    )

    print("\nAvailable Columns:\n")

    for col in first_df.columns:
        print(f"- {col}")

    user_input = input(
        "\nEnter columns for duplicate "
        "check (comma separated): "
    )

    columns = [
        col.strip()
        for col in user_input.split(",")
        if col.strip()
    ]

    for file_name, df in file_data.items():

        try:

            print(
                f"\nProcessing: {file_name}"
            )

            validate_columns(
                df,
                columns
            )

            duplicates_df, unique_df = (
                detect_duplicates(
                    df,
                    columns
                )
            )

            export_results(
                output_folder,
                file_name,
                duplicates_df,
                unique_df
            )

            print(
                f"Duplicates: "
                f"{len(duplicates_df)}"
            )

            print(
                f"Unique: "
                f"{len(unique_df)}"
            )

        except Exception as e:

            print(
                f"Skipped {file_name}: {e}"
            )

    print(
        "\nProcessing Completed Successfully."
    )


if __name__ == "__main__":
    main()
